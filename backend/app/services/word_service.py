"""
Word document generation service for analysis reports.
Generates formatted Word documents with text and chart images.
"""

import io
import re
from datetime import datetime
from typing import Optional, List, Dict
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordService:
    """Service for generating Word documents from analysis data."""

    def __init__(self):
        self.doc = None

    def generate_report(
        self, content: str, filename: Optional[str] = None, include_charts: bool = True
    ) -> bytes:
        """Generate a Word document report from analysis content."""
        self.doc = Document()

        # Add title
        title = self.doc.add_heading("Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}"
        ).italic = True

        if filename:
            source = self.doc.add_paragraph()
            source.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source.add_run(f"Source: {filename}").italic = True

        self.doc.add_paragraph()

        # Parse and add content
        self._parse_and_add_content(content)

        # Extract and add charts
        if include_charts and content:
            charts = self._extract_all_chart_data(content)
            if charts:
                self.doc.add_heading("Visualizations", level=1)
                for chart_info in charts:
                    self._add_chart_image(chart_info)

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _extract_all_chart_data(self, content: str) -> List[Dict]:
        """Extract all chart data using comprehensive patterns."""
        charts = []

        # Extract Issue Type data (Improvement, Bug, Change Request, Task)
        type_data = self._extract_generic_data(
            content,
            [
                "Improvement",
                "Bug",
                "Change Request",
                "Task",
                "Feature",
                "Epic",
                "Story",
                "Sub-task",
            ],
        )
        if type_data:
            charts.append(
                {
                    "type": "bar",
                    "labels": list(type_data.keys()),
                    "values": list(type_data.values()),
                    "title": "Issues by Type",
                }
            )

        # Extract Status data
        status_data = self._extract_generic_data(
            content,
            [
                "Open",
                "In Progress",
                "Closed",
                "Reopened",
                "Resolved",
                "Done",
                "To Do",
                "Blocked",
                "In Review",
            ],
        )
        if status_data:
            charts.append(
                {
                    "type": "pie",
                    "labels": list(status_data.keys()),
                    "values": list(status_data.values()),
                    "title": "Issues by Status",
                }
            )

        # Extract Priority data
        priority_data = self._extract_generic_data(
            content,
            [
                "Critical",
                "High",
                "Medium",
                "Low",
                "Blocker",
                "Major",
                "Minor",
                "Trivial",
            ],
        )
        if priority_data:
            charts.append(
                {
                    "type": "bar_horizontal",
                    "labels": list(priority_data.keys()),
                    "values": list(priority_data.values()),
                    "title": "Issues by Priority",
                }
            )

        return charts

    def _extract_generic_data(
        self, content: str, keywords: List[str]
    ) -> Dict[str, int]:
        """Extract data for given keywords from content using multiple patterns."""
        data = {}

        for keyword in keywords:
            # Pattern 1: "Keyword: 12" or "- Keyword: 12" or "• Keyword: 12"
            pattern1 = rf"[-•*]?\s*{re.escape(keyword)}\s*[:\-–]\s*(\d+)"
            matches = re.findall(pattern1, content, re.IGNORECASE)
            if matches:
                data[keyword] = int(matches[0])
                continue

            # Pattern 2: "12 Keyword" or "12 keyword issues"
            pattern2 = rf"(\d+)\s+{re.escape(keyword)}(?:\s+issues?)?"
            matches = re.findall(pattern2, content, re.IGNORECASE)
            if matches:
                data[keyword] = int(matches[0])
                continue

            # Pattern 3: "Keyword (12)" or "Keyword[12]"
            pattern3 = rf"{re.escape(keyword)}\s*[\(\[]\s*(\d+)\s*[\)\]]"
            matches = re.findall(pattern3, content, re.IGNORECASE)
            if matches:
                data[keyword] = int(matches[0])
                continue

            # Pattern 4: Table cell pattern "| Keyword | 12 |" or "| Keyword | value | 12 |"
            pattern4 = rf"\|\s*{re.escape(keyword)}\s*\|[^|]*?(\d+)\s*\|"
            matches = re.findall(pattern4, content, re.IGNORECASE)
            if matches:
                data[keyword] = int(
                    matches[-1]
                )  # Take last number (usually the total/count)
                continue

            # Pattern 5: "Keyword ... 12" (for data tables with spacing)
            pattern5 = rf"{re.escape(keyword)}[^\d]*?(\d+)(?:\s*$|\s*\n)"
            matches = re.findall(pattern5, content, re.IGNORECASE | re.MULTILINE)
            if matches:
                data[keyword] = int(matches[0])
                continue

            # Pattern 6: "Keyword 38.1%" - percentage (use for proportional data)
            pattern6 = rf"{re.escape(keyword)}\s*(\d+(?:\.\d+)?)\s*%"
            matches = re.findall(pattern6, content, re.IGNORECASE)
            if matches:
                data[keyword] = float(matches[0])
                continue

        return data

    def _parse_and_add_content(self, content: str):
        """Parse and add formatted content."""
        if not content:
            self.doc.add_paragraph("No content available.")
            return

        lines = content.split("\n")
        current_list = []
        in_table = False
        table_rows = []

        for line in lines:
            if "|" in line and line.strip().startswith("|"):
                if not in_table:
                    self._flush_list(current_list)
                    current_list = []
                    in_table = True
                    table_rows = []

                if not re.match(r"^\s*\|[\s\-:|]+\|\s*$", line):
                    cells = [cell.strip() for cell in line.split("|")[1:-1]]
                    if cells:
                        table_rows.append(cells)
                continue
            elif in_table:
                self._add_table(table_rows)
                table_rows = []
                in_table = False

            if line.startswith("### "):
                self._flush_list(current_list)
                current_list = []
                self.doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                self._flush_list(current_list)
                current_list = []
                self.doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                self._flush_list(current_list)
                current_list = []
                self.doc.add_heading(line[2:].strip(), level=1)
            elif line.strip().startswith(("- ", "* ", "• ")):
                current_list.append(line.strip()[2:])
            elif re.match(r"^\d+\.\s", line.strip()):
                match = re.match(r"^\d+\.\s(.+)", line.strip())
                if match:
                    current_list.append(match.group(1))
            elif line.strip():
                self._flush_list(current_list)
                current_list = []
                self._add_formatted_paragraph(line)
            elif not line.strip() and current_list:
                self._flush_list(current_list)
                current_list = []

        if in_table and table_rows:
            self._add_table(table_rows)
        self._flush_list(current_list)

    def _add_table(self, rows: List[List[str]]):
        """Add a table to the document."""
        if not rows:
            return

        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.text = cell_text
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        self.doc.add_paragraph()

    def _flush_list(self, items: List[str]):
        """Add list items."""
        for item in items:
            para = self.doc.add_paragraph(style="List Bullet")
            self._add_formatted_text(para, item)

    def _add_formatted_paragraph(self, text: str):
        """Add a formatted paragraph."""
        para = self.doc.add_paragraph()
        self._add_formatted_text(para, text)

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with markdown formatting."""
        pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))"

        for match in re.finditer(pattern, text):
            full_match = match.group(0)
            if full_match.startswith("**") and full_match.endswith("**"):
                run = paragraph.add_run(match.group(2))
                run.bold = True
            elif full_match.startswith("*") and full_match.endswith("*"):
                run = paragraph.add_run(match.group(3))
                run.italic = True
            elif full_match.startswith("`") and full_match.endswith("`"):
                run = paragraph.add_run(match.group(4))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            else:
                paragraph.add_run(match.group(5) or full_match)

    def _add_chart_image(self, chart_data: Dict):
        """Generate a chart and add to document."""
        try:
            colors = [
                "#e53e3e",
                "#ed8936",
                "#38a169",
                "#3182ce",
                "#805ad5",
                "#d69e2e",
                "#319795",
                "#dd6b20",
            ]

            if chart_data["type"] == "pie":
                fig, ax = plt.subplots(figsize=(8, 6))
                wedges, texts, autotexts = ax.pie(
                    chart_data["values"],
                    labels=chart_data["labels"],
                    autopct="%1.1f%%",
                    colors=colors[: len(chart_data["labels"])],
                    startangle=90,
                    explode=[0.03] * len(chart_data["labels"]),
                    shadow=True,
                )
                for autotext in autotexts:
                    autotext.set_fontsize(10)
                    autotext.set_fontweight("bold")
                ax.set_title(
                    chart_data.get("title", "Chart"),
                    fontsize=14,
                    fontweight="bold",
                    pad=15,
                )

            elif chart_data["type"] == "bar":
                fig, ax = plt.subplots(figsize=(10, 6))
                x = range(len(chart_data["labels"]))
                bars = ax.bar(
                    x,
                    chart_data["values"],
                    color=colors[: len(chart_data["labels"])],
                    edgecolor="white",
                )
                ax.set_xticks(x)
                ax.set_xticklabels(chart_data["labels"], fontsize=11)
                ax.set_title(
                    chart_data.get("title", "Chart"),
                    fontsize=14,
                    fontweight="bold",
                    pad=15,
                )
                ax.set_ylabel("Count", fontsize=11)
                for bar, val in zip(bars, chart_data["values"]):
                    height = bar.get_height()
                    ax.text(
                        bar.get_x() + bar.get_width() / 2.0,
                        height + 0.3,
                        f"{int(val)}",
                        ha="center",
                        va="bottom",
                        fontsize=11,
                        fontweight="bold",
                    )
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)

            elif chart_data["type"] == "bar_horizontal":
                fig, ax = plt.subplots(figsize=(10, 6))
                y = range(len(chart_data["labels"]))
                bars = ax.barh(
                    y,
                    chart_data["values"],
                    color=colors[: len(chart_data["labels"])],
                    edgecolor="white",
                )
                ax.set_yticks(y)
                ax.set_yticklabels(chart_data["labels"], fontsize=11)
                ax.set_title(
                    chart_data.get("title", "Chart"),
                    fontsize=14,
                    fontweight="bold",
                    pad=15,
                )
                ax.set_xlabel("Count", fontsize=11)
                for bar, val in zip(bars, chart_data["values"]):
                    width = bar.get_width()
                    ax.text(
                        width + 0.3,
                        bar.get_y() + bar.get_height() / 2.0,
                        f"{int(val)}",
                        ha="left",
                        va="center",
                        fontsize=11,
                        fontweight="bold",
                    )
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)
            else:
                return

            plt.tight_layout()

            img_buffer = io.BytesIO()
            plt.savefig(
                img_buffer,
                format="png",
                dpi=150,
                bbox_inches="tight",
                facecolor="white",
            )
            plt.close(fig)
            img_buffer.seek(0)

            self.doc.add_paragraph()
            self.doc.add_picture(img_buffer, width=Inches(5.5))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            caption = self.doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(chart_data.get("title", "Chart"))
            run.italic = True
            run.font.size = Pt(10)

        except Exception as e:
            para = self.doc.add_paragraph()
            para.add_run(f"[Chart error: {str(e)}]").italic = True
